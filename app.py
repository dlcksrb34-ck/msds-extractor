"""
MSDS 구성성분 추출기 - 웹 버전
Flask 기반 웹 서버
"""

import re
import os
import io
import json
from pathlib import Path
from flask import Flask, request, jsonify, render_template_string, send_file

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB 제한

# ── Tesseract 경로 설정 ───────────────────────────────────────────────────
try:
    import pytesseract
    import shutil
    # Windows
    if os.path.exists(r"C:\Program Files\Tesseract-OCR\tesseract.exe"):
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    # Linux (Railway/nixpacks) - which로 경로 자동 탐색
    elif shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = shutil.which("tesseract")
    # Linux 일반 경로
    elif os.path.exists("/usr/bin/tesseract"):
        pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"
    elif os.path.exists("/usr/local/bin/tesseract"):
        pytesseract.pytesseract.tesseract_cmd = "/usr/local/bin/tesseract"
except ImportError:
    pass

# ── 패턴 정의 (msds_gui.py와 동일) ───────────────────────────────────────
SECTION3_START = [
    re.compile(r"^3\s*[\.\)]\s*구성\s*성분", re.I),
    re.compile(r"^3\s*[\.\)]\s*성분\s*정보", re.I),
    re.compile(r"^구성\s*성분\s*의?\s*명칭\s*및\s*함유량", re.I),
    re.compile(r"^구성\s*분의?\s*명칭\s*및\s*함유량", re.I),
    re.compile(r"^3\s*[\.\_\)]\s*구성\s*분", re.I),
    re.compile(r"^3\s*[\.\_\)]\s*구성\s*및\s*성분", re.I),
    re.compile(r"^구성\s*및\s*성분", re.I),
    re.compile(r"^제\s*3\s*항\s*[\.\)]?\s*구성", re.I),
    re.compile(r"^3\s*항\s*[:\s]", re.I),
    re.compile(r"^3\s*[\.\)]\s*composition", re.I),
    re.compile(r"^3\s+composition", re.I),
    re.compile(r"^composition\s*[/\\|]\s*information\s*on\s*ingredients", re.I),
    re.compile(r"^information\s*on\s*ingredients", re.I),
    re.compile(r"^section\s*[:\-]?\s*3\b", re.I),
    re.compile(r"^3\s*[\.\)\s]\s*hazardous\s*ingredients", re.I),
    re.compile(r"^3\.\s*(homogeneous|mixture|substance|component)", re.I),
    re.compile(r"^composition\s*/\s*information\s*on\s*ingredients$", re.I),
    re.compile(r"^3\.[A-Za-z\s]{2,}(?:ient|tion|osi|pos|comp|ingr)", re.I),
]
SECTION4_START = [
    re.compile(r"^4\s*[\.\)]\s*", re.I),
    re.compile(r"^4\s*[\.\)]\s*응급\s*조치", re.I),
    re.compile(r"^4\s*[\.\)]\s*first[\s\-]*aid", re.I),
    re.compile(r"^section\s*[:\-]?\s*4\b", re.I),
    re.compile(r"^제\s*4\s*항", re.I),
    re.compile(r"^4\s*항\s*[:\s]", re.I),
]
PAGE_HEADER = re.compile(
    r"^(물질안전보건자료|safety data sheet|sds\b|material safety|"
    r"버전\s*최종\s*개정일자|version|msds\s*번호|sds\s*번호|"
    r"\d+\s*/\s*\d+\s*$)", re.I
)
CAS_RE = re.compile(r"\b(\d{2,7}-\d{1,3}-\d+)\b")
CONTENT_RE = re.compile(
    r"([\u2265\u2264≥≤]?>=?\s*\d+\.?\d*\s*[-~]\s*<?=?\s*\d+\.?\d*"
    r"|[\u2265\u2264≥≤]?>=?\s*\d+\.?\d*|[\u2265\u2264≥≤]?<=?\s*\d+\.?\d*"
    r"|[≥≤]\s*\d+\.?\d*|\d+\.?\d*\s*[-~]\s*\d+\.?\d*|\d+\.?\d*)"
    r"\s*(%\s*(?:w/w|v/v|w/v)?)?", re.I
)
INLINE_CONTENT_RE = re.compile(
    r"[(|\[]\s*~?\s*(>=?\s*\d+\.?\d*\s*[-~]\s*<?=?\s*\d+\.?\d*"
    r"|>=?\s*\d+\.?\d*|<=?\s*\d+\.?\d*|\d+\.?\d*)\s*%?\s*[)\]]"
)
KV_KEYS = {
    "chemical name": "name", "chemical names": "name", "chemical identity": "name",
    "화학물질명": "name", "화학명": "name",
    "common name": "common", "common names": "common", "관용명": "common",
    "cas number": "cas", "cas no": "cas", "cas no.": "cas",
    "cas번호": "cas", "cas 번호": "cas",
    "cas number and other unique identifiers": "cas", "cas-nr": "cas",
    "ec number": "ec", "ec no": "ec", "ec no.": "ec", "ec번호": "ec", "ec 번호": "ec",
    "concentration": "content", "content": "content",
    "% weight": "content", "% [weight]": "content",
    "함유량": "content", "함량": "content",
    "함량 (w/w)": "content", "함량 (v/v)": "content", "함량 (w/v)": "content",
    "농도": "content", "기존화학물질번호": "ke", "구성성분번호": "ke",
}

def is_section3(t): return any(p.search(t) for p in SECTION3_START)
def is_section4(t): return any(p.search(t) for p in SECTION4_START)
def is_header(t):   return bool(PAGE_HEADER.match(t))

def is_cas(s):
    parts = s.split("-")
    if not (len(parts) == 3 and len(parts[1]) == 2): return False
    digits = parts[0] + parts[1] + parts[2]
    if not digits.isdigit(): return False
    check = int(digits[-1])
    total = sum(int(d) * (i + 1) for i, d in enumerate(reversed(digits[:-1])))
    return (total % 10) == check

def fix_cas_ocr(cas):
    if is_cas(cas): return cas
    parts = cas.split("-")
    if len(parts) == 3 and len(parts[1]) == 2 and len(parts[2]) > 1:
        for trim in range(1, min(3, len(parts[2]))):
            candidate = f"{parts[0]}-{parts[1]}-{parts[2][:-trim]}"
            if is_cas(candidate): return candidate
    return cas

def clean(text):
    lines = [l.rstrip() for l in text.splitlines()]
    result, blank = [], 0
    for l in lines:
        if l == "":
            blank += 1
            if blank <= 1: result.append("")
        else:
            blank = 0; result.append(l)
    return "\n".join(result).strip()

def remove_page_headers(text):
    return "\n".join(line for line in text.splitlines()
                     if not (line.strip() and is_header(line.strip())))

def dedup_korean_chars(text):
    return re.sub(r"([\uAC00-\uD7A3\u3130-\u318F])\1{1,}", r"\1", text)

def merge_split_lines(text):
    text = re.sub(r"(\d+-)[ \t]*\n[ \t]*(\d)", r"\1\2", text)
    lines = text.splitlines()
    result = []
    i = 0
    CAS_FRONT = re.compile(r"(?<!\w)(\d{3,7})-(?!\d)")
    CAS_TAIL  = re.compile(r"\b(\d{1,2}-\d+)\b")
    while i < len(lines):
        line = lines[i]
        front_m = CAS_FRONT.search(line)
        if front_m and i + 1 < len(lines):
            next_line = lines[i + 1]
            tail_m = CAS_TAIL.search(next_line)
            if tail_m:
                cas_tail = tail_m.group(1)
                line = line[:front_m.end()] + cas_tail + " " + line[front_m.end():]
                rest = (next_line[:tail_m.start()] + next_line[tail_m.end():]).strip()
                if rest: result.append(line); result.append(rest)
                else: result.append(line)
                i += 2; continue
        result.append(line); i += 1
    text = "\n".join(result)
    text = re.sub(r"(,)[ \t]*\n[ \t]*([a-zA-Z])", r"\1 \2", text)
    return text

def slice_section3(full_text):
    text = remove_page_headers(full_text)
    lines = text.splitlines()
    collecting = False
    result_lines = []
    for line in lines:
        stripped = line.strip()
        if not collecting:
            if is_section3(stripped): collecting = True; result_lines.append(line)
        else:
            if is_section4(stripped): break
            result_lines.append(line)
    return clean("\n".join(result_lines)) if result_lines else None

def parse_kv_block(lines):
    kv = {}
    last_field = None
    for line in lines:
        stripped = line.strip()
        if not stripped: last_field = None; continue
        m_with_val = re.match(r"^([^:]{2,60}?)\s*:\s*(.+)$", stripped)
        m_no_val   = re.match(r"^([^:]{2,60}?)\s*:\s*$", stripped)
        if m_with_val:
            key_raw = re.sub(r"^\d+(\.\d+)*\.?\s*", "", m_with_val.group(1).strip().lower().rstrip(".")).strip()
            val = m_with_val.group(2).strip()
            field = KV_KEYS.get(key_raw)
            if field and field not in kv: kv[field] = val
            last_field = None
        elif m_no_val:
            key_raw = re.sub(r"^\d+(\.\d+)*\.?\s*", "", m_no_val.group(1).strip().lower().rstrip(".")).strip()
            field = KV_KEYS.get(key_raw)
            if field and field not in kv: kv[field] = ""; last_field = field
            else: last_field = None
        else:
            if stripped:
                cas_m_check = CAS_RE.search(stripped)
                if cas_m_check and is_cas(fix_cas_ocr(cas_m_check.group(1))):
                    if "cas" not in kv or not kv["cas"]: kv["cas"] = stripped
                elif last_field and not kv.get(last_field):
                    kv[last_field] = stripped
                elif not last_field and "name" not in kv:
                    if not re.match(r"^[\(\[]", stripped) and not re.match(r"^\d", stripped):
                        kv["name"] = stripped
            last_field = None

    if not kv.get("cas"): return None
    cas = kv.get("cas", "").strip()
    cas_m = CAS_RE.search(cas)
    if not cas_m or not is_cas(fix_cas_ocr(cas_m.group(1))): return None
    cas = fix_cas_ocr(cas_m.group(1))

    name = ""
    for field in ("name", "common", "synonym"):
        val = kv.get(field, "").strip()
        if not val: continue
        if re.match(r"^\d+(\.\d+)*\.?\s*", val) and len(val) < 5: continue
        name = val; break

    content = kv.get("content", "").strip()
    if not content and name:
        ic = INLINE_CONTENT_RE.search(name)
        if ic: content = ic.group(1).strip() + "%"

    return {"name": name, "cas": cas, "content": content if content else "-"}

def parse_ingredients(section_text):
    if not section_text: return []
    lines = section_text.splitlines()
    rows = []
    seen_cas = set()

    def _kv_key(line):
        m = re.match(r"^([^:]{2,60}?)\s*:\s*", line.strip())
        if not m: return None
        k = re.sub(r"^\d+(\.\d+)*\.?\s*", "", m.group(1).strip().lower().rstrip(".")).strip()
        return k

    kv_count = sum(1 for l in lines if _kv_key(l) in KV_KEYS)
    use_kv = kv_count >= 2

    if use_kv:
        block_starts = [
            i for i, l in enumerate(lines)
            if re.match(r"^(\d+(\.\d+)*\.?\s*)?(cas\s*(?:number|no\.?|number\s*and\s*other[^:]*)|cas번호|cas\s*번호)\s*:", l.strip(), re.I)
        ]
        if not block_starts:
            for i, l in enumerate(lines):
                s = l.strip()
                cas_m = CAS_RE.search(s)
                if cas_m and is_cas(fix_cas_ocr(cas_m.group(1))) and not re.match(r"^([^:]{2,60}?)\s*:", s):
                    block_starts = [i]; break
        for bs in block_starts:
            block = lines[max(0, bs - 8): bs + 6]
            result = parse_kv_block(block)
            if result and result["cas"] not in seen_cas:
                seen_cas.add(result["cas"]); rows.append(result)
        if rows: return rows

    for i, line in enumerate(lines):
        stripped = line.strip()
        all_matches = list(CAS_RE.finditer(stripped))
        cas_match = next((m for m in all_matches if is_cas(fix_cas_ocr(m.group(1)))), None)
        if not cas_match: continue
        cas = fix_cas_ocr(cas_match.group(1))
        if cas in seen_cas: continue
        seen_cas.add(cas)

        name_part = stripped[:cas_match.start()].strip(" ,-\t:")
        name_part = re.sub(r'\d{2,7}-\d{3}-\d+\s*/\s*', '', name_part).strip(" ,-/\t:")
        name_part = re.sub(r",?\s*(CAS[-\s]?No\.?|CAS\s*번호|CAS\s*number)\s*$", "", name_part, flags=re.I).strip(" ,-")
        name_part = re.sub(r"[\s,]*\b(자료\s*없음|해당\s*없음|없음|N/?A|not\s+available|unknown|none)\b[\s,]*", "", name_part, flags=re.I).strip(" ,-")
        name_part = re.sub(r",?\s*(CAS[-\s]?No\.?|CAS\s*번호|CAS\s*number)\s*$", "", name_part, flags=re.I).strip(" ,-")
        prefix_m = re.match(r"^[a-zA-Z\uAC00-\uD7A3\s\-]+:\s*", name_part)
        if prefix_m: name_part = name_part[prefix_m.end():].strip()
        if len(name_part) < 4 and i > 0:
            prev = lines[i - 1].strip()
            if prev and not CAS_RE.search(prev) and not is_header(prev):
                name_part = (prev + " " + name_part).strip()

        after_cas = stripped[cas_match.end():].strip()
        after_cas = re.sub(r"^\d+(\.\d+)+\.?\s+", "", after_cas).strip()
        after_cas_clean = re.sub(r"[/\s]*[A-Z]{1,4}-\d{4,}[/\s]*", " ", after_cas).strip()
        after_cas_clean = re.sub(r"\b\d{3}-\d{3}-\d{1,2}\b", " ", after_cas_clean).strip()

        content = ""
        for cm in CONTENT_RE.finditer(after_cas_clean):
            val = cm.group(0).strip()
            if not val: continue
            num_m = re.search(r"\d+\.?\d*", val)
            if num_m and float(num_m.group()) > 100.0: continue
            content = val; break

        if not content and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if not re.match(r"^\d+(\.\d+)*\.?\s", next_line):
                for cm2 in CONTENT_RE.finditer(next_line):
                    val = cm2.group(0).strip()
                    if not val: continue
                    num_m2 = re.search(r"\d+\.?\d*", val)
                    if num_m2 and float(num_m2.group()) <= 100.0:
                        content = val; break

        if name_part:
            rows.append({"name": name_part, "cas": cas, "content": content if content else "-"})

    return rows

def extract_from_pdf_text(path):
    import pdfplumber
    full_text = ""
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text(x_tolerance=2, y_tolerance=3)
            if t: full_text += t + "\n"
    return slice_section3(merge_split_lines(dedup_korean_chars(full_text)))

def extract_from_pdf_ocr(path, lang="kor+eng"):
    from pdf2image import convert_from_path
    import pytesseract, shutil
    # 경로 재확인
    if shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = shutil.which("tesseract")
    images = convert_from_path(str(path), dpi=300)
    full_text = ""
    for img in images:
        try:
            full_text += pytesseract.image_to_string(img, lang=lang) + "\n"
        except Exception:
            full_text += pytesseract.image_to_string(img, lang="eng") + "\n"
    return slice_section3(merge_split_lines(dedup_korean_chars(full_text)))

def extract_from_pdf(path):
    import pdfplumber
    sample = ""
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages[:5]:
            t = page.extract_text()
            if t: sample += t
    if len(sample.strip()) > 100:
        result = extract_from_pdf_text(path)
        if result: return result, "텍스트형 PDF"
    result = extract_from_pdf_ocr(path)
    return result, "스캔 PDF (OCR)"

def extract_from_docx(path):
    from docx import Document
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            full_text += "\n" + "\t".join(cell.text for cell in row.cells)
    return slice_section3(merge_split_lines(dedup_korean_chars(full_text))), "Word 문서"

def extract_section3(path):
    ext = path.suffix.lower()
    if ext == ".pdf": return extract_from_pdf(path)
    elif ext in (".docx", ".doc"): return extract_from_docx(path)
    return None, "지원하지 않는 형식"


# ── HTML 템플릿 ───────────────────────────────────────────────────────────
HTML = open("templates/index.html", encoding="utf-8").read()


# ── 라우트 ────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template_string(HTML)

@app.route("/extract", methods=["POST"])
def extract():
    if "file" not in request.files:
        return jsonify({"error": "파일이 없습니다."}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "파일명이 없습니다."}), 400

    suffix = Path(file.filename).suffix.lower()
    if suffix not in (".pdf", ".docx", ".doc"):
        return jsonify({"error": "PDF 또는 Word 파일만 지원합니다."}), 400

    tmp_path = Path(f"/tmp/{file.filename}")
    file.save(str(tmp_path))

    try:
        text, mode = extract_section3(tmp_path)
        ingredients = parse_ingredients(text) if text else []
        return jsonify({
            "success": True,
            "filename": file.filename,
            "mode": mode,
            "text": text or "",
            "ingredients": ingredients,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        if tmp_path.exists(): tmp_path.unlink()

@app.route("/download/excel", methods=["POST"])
def download_excel():
    data = request.get_json()
    ingredients = data.get("ingredients", [])
    filename = data.get("filename", "result")

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "구성성분"

        hdr_font  = Font(bold=True, color="FFFFFF", size=10)
        hdr_fill  = PatternFill("solid", fgColor="0F3460")
        even_fill = PatternFill("solid", fgColor="E8EEF7")
        thin      = Side(style="thin", color="BEC8D9")
        border    = Border(left=thin, right=thin, top=thin, bottom=thin)
        center    = Alignment(horizontal="center", vertical="center")
        left_a    = Alignment(horizontal="left",   vertical="center")

        headers    = ["No.", "화학물질명", "CAS 번호", "함유량 (% w/w)", "파일명"]
        col_widths = [6, 55, 16, 18, 40]
        for col, (h, w) in enumerate(zip(headers, col_widths), 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = hdr_font; cell.fill = hdr_fill
            cell.alignment = center; cell.border = border
            ws.column_dimensions[cell.column_letter].width = w

        for i, ing in enumerate(ingredients, 1):
            fill = even_fill if i % 2 == 0 else PatternFill()
            vals   = [i, ing["name"], ing["cas"], ing["content"], filename]
            aligns = [center, left_a, center, center, left_a]
            for col, (v, al) in enumerate(zip(vals, aligns), 1):
                cell = ws.cell(row=i+1, column=col, value=v)
                cell.alignment = al; cell.border = border
                if fill.fill_type: cell.fill = fill

        ws.auto_filter.ref = f"A1:E{len(ingredients)+1}"
        ws.freeze_panes = "A2"

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return send_file(output, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                         as_attachment=True, download_name=f"{Path(filename).stem}_section3.xlsx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/download/excel/all", methods=["POST"])
def download_excel_all():
    data = request.get_json()
    all_results = data.get("results", [])

    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "구성성분_전체"

        hdr_font  = Font(bold=True, color="FFFFFF", size=10)
        hdr_fill  = PatternFill("solid", fgColor="0F3460")
        even_fill = PatternFill("solid", fgColor="E8EEF7")
        thin      = Side(style="thin", color="BEC8D9")
        border    = Border(left=thin, right=thin, top=thin, bottom=thin)
        center    = Alignment(horizontal="center", vertical="center")
        left_a    = Alignment(horizontal="left",   vertical="center")

        headers    = ["No.", "화학물질명", "CAS 번호", "함유량 (% w/w)", "파일명"]
        col_widths = [6, 55, 16, 18, 40]
        for col, (h, w) in enumerate(zip(headers, col_widths), 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = hdr_font; cell.fill = hdr_fill
            cell.alignment = center; cell.border = border
            ws.column_dimensions[cell.column_letter].width = w

        row_idx = 2
        global_no = 1
        for entry in all_results:
            filename    = entry.get("filename", "")
            ingredients = entry.get("ingredients", [])
            for ing in ingredients:
                fill = even_fill if global_no % 2 == 0 else PatternFill()
                vals   = [global_no, ing["name"], ing["cas"], ing["content"], filename]
                aligns = [center, left_a, center, center, left_a]
                for col, (v, al) in enumerate(zip(vals, aligns), 1):
                    cell = ws.cell(row=row_idx, column=col, value=v)
                    cell.alignment = al; cell.border = border
                    if fill.fill_type: cell.fill = fill
                row_idx += 1
                global_no += 1

        ws.auto_filter.ref = f"A1:E{row_idx - 1}"
        ws.freeze_panes = "A2"

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return send_file(output,
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                         as_attachment=True,
                         download_name="MSDS_전체_구성성분.xlsx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
